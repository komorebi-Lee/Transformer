import os
import logging
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from datetime import datetime

logger = logging.getLogger(__name__)


class WordExporter:
    """Word文档导出器 - 支持超链接和导航"""

    def __init__(self):
        self.document = None

    def export_structured_codes_with_hyperlinks(self, file_path: str,
                                                structured_codes: Dict[str, Any],
                                                combined_text: str,
                                                text_mapping: Dict[str, Any]) -> bool:
        """导出编码结构到Word文档，包含超链接"""
        try:
            self.document = Document()

            # 设置文档属性
            self.document.core_properties.title = "扎根理论编码分析结果"
            self.document.core_properties.author = "扎根理论编码分析系统"
            self.document.core_properties.comments = "自动生成的扎根理论三级编码分析报告"

            # 添加标题
            title = self.document.add_heading('扎根理论编码分析结果', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加导出信息
            self.document.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            self.document.add_paragraph(f"三阶编码数量: {len(structured_codes)}")

            total_second = sum(len(categories) for categories in structured_codes.values())
            total_first = sum(len(contents) for categories in structured_codes.values()
                              for contents in categories.values())
            self.document.add_paragraph(f"二阶编码数量: {total_second}")
            self.document.add_paragraph(f"一阶编码数量: {total_first}")

            self.document.add_paragraph()  # 空行

            # 第一部分：三级编码表格
            self.add_coding_table(structured_codes)

            # 分页
            self.document.add_page_break()

            # 第二部分：合并文本与引用
            self.add_combined_text_with_references(combined_text, structured_codes, text_mapping)

            # 保存文档
            self.document.save(file_path)
            logger.info(f"Word文档已导出: {file_path}")
            return True

        except Exception as e:
            logger.error(f"导出Word文档失败: {e}")
            return False

    def add_coding_table(self, structured_codes: Dict[str, Any]):
        """添加编码表格"""
        # 添加标题
        heading = self.document.add_heading('一、三级编码结构', level=1)

        # 创建表格
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'

        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '三阶编码'
        hdr_cells[1].text = '二阶编码'
        hdr_cells[2].text = '一阶编码'

        # 设置表头格式
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True

        # 填充数据
        for third_category, second_categories in structured_codes.items():
            for second_category, first_contents in second_categories.items():
                for content_data in first_contents:
                    # 添加新行
                    row_cells = table.add_row().cells

                    # 三阶编码
                    row_cells[0].text = self.clean_category_name(third_category)

                    # 二阶编码
                    row_cells[1].text = self.clean_category_name(second_category)

                    # 一阶编码
                    if isinstance(content_data, dict):
                        content = content_data.get('content', '')
                        code_id = content_data.get('code_id', '')
                    else:
                        content = str(content_data)
                        code_id = ""

                    clean_content = self.clean_first_level_content(content)
                    row_cells[2].text = clean_content

                    # 如果有编码ID，添加书签
                    if code_id:
                        self.add_bookmark_to_paragraph(row_cells[2].paragraphs[0], code_id)

    def add_combined_text_with_references(self, combined_text: str,
                                          structured_codes: Dict[str, Any],
                                          text_mapping: Dict[str, Any]):
        """添加合并文本和引用"""
        # 添加标题
        heading = self.document.add_heading('二、原始文本与编码引用', level=1)

        # 添加说明
        self.document.add_paragraph("以下为所有访谈文本的合并内容，其中标记的文本对应上方的一阶编码。")

        # 处理文本，添加引用标记
        text_paragraph = self.document.add_paragraph()

        # 收集所有编码位置
        code_positions = self.extract_code_positions(structured_codes, text_mapping)

        # 按位置排序
        sorted_positions = sorted(code_positions.items(), key=lambda x: x[1]['start'])

        current_pos = 0
        for code_id, position_info in sorted_positions:
            start_pos = position_info['start']
            end_pos = position_info['end']
            content = position_info['content']

            # 添加位置之前的文本
            if current_pos < start_pos:
                text_before = combined_text[current_pos:start_pos]
                text_paragraph.add_run(text_before)

            # 添加带标记的编码文本
            coded_text = combined_text[start_pos:end_pos]
            run = text_paragraph.add_run(coded_text)
            run.bold = True
            run.font.color.rgb = (255, 0, 0)  # 红色

            # 添加上标引用
            ref_run = text_paragraph.add_run(f'[{code_id}]')
            ref_run.font.superscript = True
            ref_run.font.color.rgb = (0, 0, 255)  # 蓝色

            current_pos = end_pos

        # 添加剩余文本
        if current_pos < len(combined_text):
            text_remaining = combined_text[current_pos:]
            text_paragraph.add_run(text_remaining)

    def extract_code_positions(self, structured_codes: Dict[str, Any],
                               text_mapping: Dict[str, Any]) -> Dict[str, Any]:
        """提取编码位置信息"""
        code_positions = {}

        for third_category, second_categories in structured_codes.items():
            for second_category, first_contents in second_categories.items():
                for content_data in first_contents:
                    if isinstance(content_data, dict):
                        code_id = content_data.get('code_id', '')
                        sentence_details = content_data.get('sentence_details', [])

                        if code_id and sentence_details:
                            # 使用第一个句子的位置
                            first_sentence = sentence_details[0]
                            content = first_sentence.get('content', '')

                            if content:
                                code_positions[code_id] = {
                                    'content': content,
                                    'start': first_sentence.get('start_position', 0),
                                    'end': first_sentence.get('end_position', len(content)),
                                    'filename': first_sentence.get('filename', '')
                                }

        return code_positions

    def add_bookmark_to_paragraph(self, paragraph, bookmark_name: str):
        """添加书签到段落"""
        # 创建书签开始
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), '1')
        bookmark_start.set(qn('w:name'), bookmark_name)
        paragraph._p.append(bookmark_start)

        # 创建书签结束
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), '1')
        bookmark_end.set(qn('w:name'), bookmark_name)
        paragraph._p.append(bookmark_end)

    def clean_category_name(self, name: str) -> str:
        """清理类别名称"""
        import re
        return re.sub(r'^[A-Z]\d*\s*', '', name.strip())

    def clean_first_level_content(self, content: str) -> str:
        """清理一阶编码内容"""
        import re
        return re.sub(r'^[A-Z]\d+\s*', '', content.strip())