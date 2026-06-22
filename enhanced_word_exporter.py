import os
import logging
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from datetime import datetime

logger = logging.getLogger(__name__)


class EnhancedWordExporter:
    """增强的Word文档导出器 - 支持新的编号系统和超链接"""

    def __init__(self):
        self.document = None

    def export_structured_codes_with_hyperlinks(self, file_path: str,
                                                structured_codes: Dict[str, Any],
                                                combined_text: str,
                                                text_mapping: Dict[str, Any],
                                                higher_level_data: list = None) -> bool:
        """导出编码结构到Word文档，包含超链接和编号系统"""
        try:
            if higher_level_data is None:
                higher_level_data = []
            self.document = Document()

            # 设置文档属性
            self.document.core_properties.title = "扎根理论编码分析结果"
            self.document.core_properties.author = "扎根理论编码分析系统"
            self.document.core_properties.comments = "自动生成的扎根理论多级编码分析报告"

            # 添加标题
            title = self.document.add_heading('扎根理论编码分析结果', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加导出信息
            self.document.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            # 统计编码数量
            total_third = len(structured_codes)
            total_second = sum(len(categories) for categories in structured_codes.values())
            total_first = sum(len(contents) for categories in structured_codes.values()
                              for contents in categories.values())

            self.document.add_paragraph(f"三阶编码数量: {total_third}")
            self.document.add_paragraph(f"二阶编码数量: {total_second}")
            self.document.add_paragraph(f"一阶编码数量: {total_first}")

            has_higher = len(higher_level_data) > 0
            if has_higher:
                self.document.add_paragraph(f"四/五/六阶编码组数量: {len(higher_level_data)}")

            self.document.add_paragraph()  # 空行

            # 根据合并文本和句子来源构建文件顺序与一阶编码显示编号映射
            file_index_map = self._build_file_index_map_from_combined_text(combined_text)
            code_display_map = self._build_first_level_display_id_map(structured_codes, file_index_map)

            # 第一部分：编码表格
            if has_higher:
                self.add_multi_order_coding_table(structured_codes, higher_level_data, code_display_map)
            else:
                self.add_coding_table_with_new_format(structured_codes, code_display_map)

            # 分页
            self.document.add_page_break()

            # 第二部分：合并文本与引用
            self.add_combined_text_with_numbered_references(combined_text, structured_codes)

            # 保存文档
            self.document.save(file_path)
            logger.info(f"Word文档已导出: {file_path}")
            return True

        except Exception as e:
            logger.error(f"导出Word文档失败: {e}")
            return False

    def add_coding_table_with_new_format(self, structured_codes: Dict[str, Any], code_display_map: Dict[str, str] = None):
        """添加编码表格 - 新编号格式"""
        # 添加标题
        heading = self.document.add_heading('一、三级编码结构', level=1)

        # 创建表格
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'

        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '三阶编码'
        hdr_cells[1].text = '二阶编码'
        hdr_cells[2].text = '一阶编码'

        # 设置表头格式
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True

        # 填充数据
        for third_category, second_categories in structured_codes.items():
            # 提取三阶编码字母（A, B, C等）
            third_letter = third_category.split(' ')[0] if ' ' in third_category else third_category[0]

            for second_category, first_contents in second_categories.items():
                # 提取二阶编码编号（A1, A2等）
                second_number = second_category.split(' ')[0] if ' ' in second_category else second_category

                for content_data in first_contents:
                    # 添加新行
                    row_cells = table.add_row().cells

                    # 三阶编码（只显示一次，避免重复）
                    row_cells[0].text = third_category

                    # 二阶编码（只显示一次，避免重复）
                    row_cells[1].text = second_category

                    # 一阶编码
                    if isinstance(content_data, dict):
                        numbered_content = content_data.get('numbered_content', '')
                        code_id = content_data.get('code_id', '')
                    else:
                        numbered_content = str(content_data)
                        code_id = ""

                    row_cells[2].text = numbered_content

                    # 如果有编码ID，添加书签
                    if code_id:
                        self.add_bookmark_to_paragraph(row_cells[2].paragraphs[0], f"code_{code_id}")

    def add_combined_text_with_numbered_references(self, combined_text: str,
                                                   structured_codes: Dict[str, Any]):
        """添加合并文本和编号引用"""
        # 添加标题
        heading = self.document.add_heading('二、原始文本与编码引用', level=1)

        # 添加说明
        self.document.add_paragraph("以下为所有访谈文本的合并内容，其中标记的文本对应上方的一阶编码。")
        self.document.add_paragraph("编号说明：")
        self.document.add_paragraph("  - A, B, C...: 三阶编码")
        self.document.add_paragraph("  - A1, A2, B1, B2...: 二阶编码")
        self.document.add_paragraph("  - A11, A12, B21, B22...: 一阶编码")

        # 处理文本，添加引用标记
        text_paragraph = self.document.add_paragraph()

        # 收集所有编码位置
        code_references = self.extract_code_references(structured_codes)

        # 按位置排序（简化处理，按在文本中出现的顺序）
        processed_text = combined_text

        # 在文本中标记编码引用
        for code_id, reference_info in code_references.items():
            original_content = reference_info.get('original_content', '')
            marked_content = reference_info.get('marked_content', '')

            if original_content and marked_content and original_content in processed_text:
                processed_text = processed_text.replace(original_content, marked_content)

        # 添加处理后的文本
        text_paragraph.add_run(processed_text)

    def extract_code_references(self, structured_codes: Dict[str, Any]) -> Dict[str, Any]:
        """提取编码引用信息"""
        code_references = {}

        for third_category, second_categories in structured_codes.items():
            for second_category, first_contents in second_categories.items():
                for content_data in first_contents:
                    if isinstance(content_data, dict):
                        code_id = content_data.get('code_id', '')
                        sentence_details = content_data.get('sentence_details', [])

                        if code_id and sentence_details:
                            # 使用第一个句子的位置
                            first_sentence = sentence_details[0]
                            original_content = first_sentence.get('original_content', '')
                            marked_content = first_sentence.get('content', original_content)

                            if original_content:
                                code_references[code_id] = {
                                    'original_content': original_content,
                                    'marked_content': marked_content,
                                    'filename': first_sentence.get('filename', '')
                                }

        return code_references

    def add_multi_order_coding_table(self, structured_codes: Dict[str, Any],
                                      higher_level_data: list,
                                      code_display_map: Dict[str, str] = None):
        """添加多阶编码表格 - 支持1-6阶完整结构"""
        if code_display_map is None:
            code_display_map = {}
        heading = self.document.add_heading('一、编码结构', level=1)

        col_count = 6
        table = self.document.add_table(rows=1, cols=col_count)
        table.style = 'Light Grid Accent 1'

        hdr_cells = table.rows[0].cells
        headers = ['六阶编码', '五阶编码', '四阶编码', '三阶编码', '二阶编码', '一阶编码']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True

        covered = set()

        for item_data in higher_level_data:
            self._add_higher_rows(table, item_data, [], code_display_map, covered)

        for third_category, second_categories in structured_codes.items():
            clean_third = self.clean_category_name(third_category)
            for second_category, first_contents in second_categories.items():
                clean_second = self.clean_category_name(second_category)
                for content_data in first_contents:
                    code_id = ''
                    display_text = ''
                    if isinstance(content_data, dict):
                        code_id = content_data.get('code_id', '')
                        base_content = content_data.get('content', '') or content_data.get('numbered_content', '')
                        clean_content = self.clean_first_level_content(base_content) if base_content else ''
                        display_id = code_display_map.get(code_id, code_id)
                        if display_id and clean_content:
                            display_text = f"{display_id} {clean_content}"
                        elif base_content:
                            display_text = base_content
                        else:
                            display_text = content_data.get('numbered_content', '') or ''
                        key = (clean_third, clean_second, code_id)
                    else:
                        display_text = str(content_data)
                        key = (clean_third, clean_second, display_text)

                    if key not in covered:
                        row_cells = table.add_row().cells
                        for c in range(3):
                            row_cells[c].text = ''
                        row_cells[3].text = third_category
                        row_cells[4].text = second_category
                        row_cells[5].text = display_text
                        if code_id:
                            self.add_bookmark_to_paragraph(row_cells[5].paragraphs[0], f"code_{code_id}")

    def _add_higher_rows(self, table, item_data: Dict[str, Any],
                         parent_path: list, code_display_map: Dict[str, str],
                         covered: set):
        """递归添加高阶编码行"""
        text = item_data.get('text', '')
        data = item_data.get('data', {}) or {}
        level = data.get('level', 3) if isinstance(data, dict) else 3

        current_path = parent_path + [text]

        children = item_data.get('children', [])
        if not children and data:
            first_items = data.get('children', []) if isinstance(data, dict) else []
            children = first_items

        if not children:
            return

        for child in children:
            child_text = child.get('text', '')
            child_data = child.get('data', {}) or {}
            child_level = child_data.get('level', 3) if isinstance(child_data, dict) else 3

            if child_level in (4, 5, 6):
                self._add_higher_rows(table, child, current_path, code_display_map, covered)
            elif child_level == 3:
                self._add_third_order_rows(table, child, current_path, code_display_map, covered)
            elif child_level == 2:
                self._add_second_order_rows_from_table(table, child, current_path, code_display_map, covered)
            elif child_level == 1:
                row_cells = table.add_row().cells
                self._fill_higher_columns(row_cells, current_path)
                display_text = self._get_first_display_text(child, code_display_map)
                row_cells[5].text = display_text

    def _add_third_order_rows(self, table, third_item: Dict, parent_path: list,
                               code_display_map: Dict, covered: set):
        """从三阶节点添加行"""
        third_text = third_item.get('text', '')
        children = third_item.get('children', [])
        for child in children:
            child_data = child.get('data', {}) or {}
            child_level = child_data.get('level', 2) if isinstance(child_data, dict) else 2
            if child_level == 2:
                second_text = child.get('text', '')
                grand_children = child.get('children', [])
                for gc in grand_children:
                    gc_data = gc.get('data', {}) or {}
                    gc_level = gc_data.get('level', 1) if isinstance(gc_data, dict) else 1
                    if gc_level == 1:
                        code_id = gc_data.get('code_id', '') if isinstance(gc_data, dict) else ''
                        display_text = self._get_first_display_text(gc, code_display_map)
                        clean_third = self.clean_category_name(third_text)
                        clean_second = self.clean_category_name(second_text)
                        key = (clean_third, clean_second, code_id)
                        covered.add(key)

                        row_cells = table.add_row().cells
                        self._fill_higher_columns(row_cells, parent_path)
                        row_cells[3].text = third_text
                        row_cells[4].text = second_text
                        row_cells[5].text = display_text
                        if code_id:
                            self.add_bookmark_to_paragraph(row_cells[5].paragraphs[0], f"code_{code_id}")

    def _add_second_order_rows_from_table(self, table, second_item: Dict, parent_path: list,
                                          code_display_map: Dict, covered: set):
        """从二阶节点添加行"""
        second_text = second_item.get('text', '')
        children = second_item.get('children', [])
        for gc in children:
            gc_data = gc.get('data', {}) or {}
            gc_level = gc_data.get('level', 1) if isinstance(gc_data, dict) else 1
            if gc_level == 1:
                code_id = gc_data.get('code_id', '') if isinstance(gc_data, dict) else ''
                display_text = self._get_first_display_text(gc, code_display_map)
                clean_third = self.clean_category_name('')
                clean_second = self.clean_category_name(second_text)
                key = (clean_third, clean_second, code_id)
                covered.add(key)

                row_cells = table.add_row().cells
                self._fill_higher_columns(row_cells, parent_path)
                row_cells[3].text = ''
                row_cells[4].text = second_text
                row_cells[5].text = display_text
                if code_id:
                    self.add_bookmark_to_paragraph(row_cells[5].paragraphs[0], f"code_{code_id}")

    def _fill_higher_columns(self, row_cells, parent_path: list):
        """填充高阶列（6/5/4阶）"""
        start_col = 6 - len(parent_path)
        for i, name in enumerate(parent_path):
            if start_col + i < 6:
                row_cells[start_col + i].text = name
        for c in range(start_col):
            row_cells[c].text = ''

    def _get_first_display_text(self, item: Dict, code_display_map: Dict) -> str:
        """获取一阶编码的显示文本"""
        data = item.get('data', {}) or {}
        text = item.get('text', '')
        if isinstance(data, dict):
            code_id = data.get('code_id', '')
            content = data.get('content', '') or data.get('numbered_content', '')
            clean_content = self.clean_first_level_content(content) if content else ''
            display_id = code_display_map.get(code_id, code_id) if code_display_map else ''
            if display_id and clean_content:
                return f"{display_id} {clean_content}"
            if content:
                return content
        return text

    def _build_file_index_map_from_combined_text(self, combined_text: str) -> Dict[str, int]:
        """根据合并文本中出现的文件顺序构建文件索引映射。

        约定：combined_text 由 main_window.get_combined_text 构造，形如
        "\n\n=== 文件名1 ===\n\n...文本...\n\n=== 文件名2 ===\n\n..."。
        """
        import re

        file_index_map: Dict[str, int] = {}
        if not combined_text:
            return file_index_map

        pattern = r"===\s*(.+?)\s*==="
        matches = re.findall(pattern, combined_text)

        index = 1
        for name in matches:
            name = name.strip()
            if name and name not in file_index_map:
                file_index_map[name] = index
                index += 1

        return file_index_map

    def _build_first_level_display_id_map(self, structured_codes: Dict[str, Any],
                                          file_index_map: Dict[str, int]) -> Dict[str, str]:
        """基于文件顺序构建一阶编码显示编号映射。

        对每个一阶编码，根据其 sentence_details 中的 filename/file_path
        找到所属文件的索引 i（从1开始），并在该文件范围内按出现顺序
        赋予本地序号 j（从1开始），最终生成显示编号 A{i}-{j:02d}。
        """
        from collections import defaultdict

        display_map: Dict[str, str] = {}

        if not structured_codes or not file_index_map:
            return display_map

        per_file_counter: Dict[int, int] = defaultdict(int)

        for _, second_categories in structured_codes.items():
            for _, first_contents in second_categories.items():
                for content_data in first_contents:
                    if not isinstance(content_data, dict):
                        continue

                    code_id = content_data.get('code_id', '')
                    if not code_id or code_id in display_map:
                        # 跳过没有编号或已处理过的编码
                        continue

                    sentence_details = content_data.get('sentence_details', [])
                    if not sentence_details:
                        continue

                    first_sentence = sentence_details[0]
                    filename = first_sentence.get('filename', '')
                    if not filename:
                        file_path = first_sentence.get('file_path', '')
                        if file_path:
                            filename = os.path.basename(file_path)

                    if not filename:
                        continue

                    file_idx = file_index_map.get(filename)
                    if not file_idx:
                        continue

                    per_file_counter[file_idx] += 1
                    local_index = per_file_counter[file_idx]
                    display_map[code_id] = f"A{file_idx}-{local_index:02d}"

        return display_map

    def add_bookmark_to_paragraph(self, paragraph, bookmark_name: str):
        """添加书签到段落"""
        try:
            # 创建书签开始
            bookmark_start = OxmlElement('w:bookmarkStart')
            bookmark_start.set(qn('w:id'), '1')
            bookmark_start.set(qn('w:name'), bookmark_name)
            paragraph._p.append(bookmark_start)

            # 创建书签结束
            bookmark_end = OxmlElement('w:bookmarkEnd')
            bookmark_end.set(qn('w:id'), '1')
            bookmark_end.set(qn('w:name'), bookmark_name)
            paragraph._p.append(bookmark_end)
        except Exception as e:
            logger.warning(f"添加书签失败: {e}")

    def clean_category_name(self, name: str) -> str:
        """清理类别名称"""
        import re
        return re.sub(r'^[A-Z]\d*\s*', '', name.strip())

    def clean_first_level_content(self, content: str) -> str:
        """清理一阶编码内容"""
        import re
        return re.sub(r'^[A-Z]\d+\s*', '', content.strip())