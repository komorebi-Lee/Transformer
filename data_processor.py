import os
import json
import re
import logging
from datetime import datetime
from typing import Dict, List, Any, Tuple, Set, Optional
import pandas as pd
from docx import Document
import copy
from collections import defaultdict

from config import Config

logger = logging.getLogger(__name__)

# 尝试导入可选依赖
try:
    import win32com.client as win32
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False

try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False

try:
    from text_numbering import TextNumberingManager
    TEXT_NUMBERING_AVAILABLE = True
except ImportError:
    TEXT_NUMBERING_AVAILABLE = False
    TextNumberingManager = None


class DataProcessor:
    """数据处理器 - 支持多文件处理和Word文档"""

    def __init__(self):
        self.sentence_counter = 0
        self.file_sentence_mapping = {}
        
        # 初始化文本编号管理器
        if TEXT_NUMBERING_AVAILABLE and TextNumberingManager:
            self.numbering_manager = TextNumberingManager()
        else:
            self.numbering_manager = None
        
        # 集成精准受访者提取（SpeakerRoleExtractor）
        try:
            from speaker_role_extractor import SpeakerRoleExtractor
            self.speaker_extractor = SpeakerRoleExtractor(use_qa_classifier=True)
            self.use_advanced_extraction = True
            logger.info("已启用精准受访者提取（SpeakerRoleExtractor + QA分类器）")
        except ImportError as e:
            self.speaker_extractor = None
            self.use_advanced_extraction = False
            logger.info(f"使用原有的段落识别（SpeakerRoleExtractor不可用: {e}）")

    def read_doc_file(self, file_path: str) -> str:
        """读取.doc文件，使用多种备选方法"""
        # 方法1: 尝试使用win32com（仅Windows）
        if WIN32COM_AVAILABLE:
            try:
                return self._read_doc_with_win32com(file_path)
            except Exception as e:
                logger.warning(f"win32com读取失败，尝试其他方法: {e}")
        
        # 方法2: 尝试使用pypandoc
        if PYPANDOC_AVAILABLE:
            try:
                import pypandoc
                return self._read_doc_with_pypandoc(file_path)
            except Exception as e:
                logger.warning(f"pypandoc读取失败，尝试其他方法: {e}")
        
        # 方法3: 尝试使用pypandoc或win32com作为降级方案
        try:
            if PYPANDOC_AVAILABLE:
                return self._read_doc_with_antiword(file_path)  # 使用更新后的函数
            elif WIN32COM_AVAILABLE:
                return self._read_doc_with_win32com(file_path)
        except Exception as e:
            logger.warning(f"降级方法读取失败: {e}")
        
        # 如果所有方法都失败，抛出异常
        raise Exception("无法读取.doc文件，请安装win32com或pypandoc")

    def _read_doc_with_win32com(self, file_path: str) -> str:
        """使用win32com读取.doc文件"""
        if not WIN32COM_AVAILABLE:
            raise Exception("win32com不可用")
        word_app = win32.Dispatch('Word.Application')
        word_app.Visible = False  # 不显示Word界面
        
        # 打开文档
        doc = word_app.Documents.Open(os.path.abspath(file_path))
        
        # 提取文本内容
        content = doc.Content.Text
        
        # 关闭文档和Word应用程序
        doc.Close()
        word_app.Quit()
        
        logger.info(f"成功使用win32com读取Word文档(.doc): {file_path}")
        return content.strip()

    def _read_doc_with_pypandoc(self, file_path: str) -> str:
        """使用pypandoc读取.doc文件"""
        # 将.doc文件转换为文本
        if not PYPANDOC_AVAILABLE:
            raise Exception("pypandoc不可用")
        content = pypandoc.convert_file(file_path, 'plain')
        logger.info(f"成功使用pypandoc读取Word文档(.doc): {file_path}")
        return content.strip()

    def _read_doc_with_antiword(self, file_path: str) -> str:
        """使用antiword命令行工具读取.doc文件"""
        import subprocess
        # 使用text-from-docx库处理.doc文件
        # 首先尝试将.doc转换为.docx临时文件，然后读取
        try:
            import zipfile
            import tempfile
            
            # 创建一个临时的docx文件，将doc内容转换
            # 这里使用pypandoc作为备选方案，如果可用
            if PYPANDOC_AVAILABLE:
                import tempfile
                import os
                
                # 创建临时文件
                temp_docx_path = None
                try:
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                        temp_docx_path = temp_file.name
                    
                    # 尝试使用pypandoc转换.doc到.docx
                    if not PYPANDOC_AVAILABLE:
                        raise Exception("pypandoc不可用")
                    pypandoc.convert_file(file_path, 'docx', outputfile=temp_docx_path)
                    
                    # 读取转换后的.docx文件
                    doc = Document(temp_docx_path)
                    content = ""
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            content += paragraph.text + "\n"
                    
                    logger.info(f"成功使用pypandoc转换并读取.doc文件: {file_path}")
                    return content.strip()
                except Exception as e:
                    logger.warning(f"pypandoc转换.doc失败: {e}")
                    # 如果转换失败，尝试其他方法
                    if temp_docx_path and os.path.exists(temp_docx_path):
                        try:
                            os.unlink(temp_docx_path)
                        except:
                            pass
                    raise e
            else:
                # 如果pypandoc不可用，尝试使用win32com作为降级方案
                if WIN32COM_AVAILABLE:
                    return self._read_doc_with_win32com(file_path)
                else:
                    raise Exception("缺少必要的依赖来读取.doc文件")
        except Exception as e:
            raise e

    def read_text_file(self, file_path: str) -> str:
        """读取文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            logger.info(f"成功读取文本文件: {file_path}")
            return content
        except Exception as e:
            logger.error(f"读取文本文件失败 {file_path}: {e}")
            raise

    def read_word_file(self, file_path: str) -> str:
        """读取Word文档，支持.doc和.docx格式"""
        try:
            # 根据文件扩展名选择适当的读取方法
            if file_path.lower().endswith('.docx'):
                # 使用python-docx读取.docx文件
                doc = Document(file_path)
                content = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content += paragraph.text + "\n"
                logger.info(f"成功读取Word文档: {file_path}")
                return content.strip()
            elif file_path.lower().endswith('.doc'):
                # 使用win32com读取.doc文件
                return self.read_doc_file(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_path}")
        except Exception as e:
            logger.error(f"读取Word文档失败 {file_path}: {e}")
            # 降级到文本读取
            return self.read_text_file(file_path)

    def read_file(self, file_path: str) -> str:
        """读取文件（自动判断类型）"""
        file_lower = file_path.lower()
        if file_lower.endswith('.docx') or file_lower.endswith('.doc'):
            return self.read_word_file(file_path)
        else:
            return self.read_text_file(file_path)

    def process_multiple_files(self, file_paths: List[str], number_mappings: Dict[str, Dict[int, str]] = None) -> Dict[str, Any]:
        """
        处理多个文件，返回统一的文本和文件映射
        
        Args:
            file_paths: 文件路径列表
            number_mappings: TextNumbering 编号映射 {filename: {number: text}}
        """
        all_texts = []
        file_sentence_mapping = {}
        self._processing_sentence_counter = 0

        for file_path in file_paths:
            try:
                # 读取文件
                content = self.read_file(file_path)
                filename = os.path.basename(file_path)
                sentence_number_lookup = self._build_sentence_number_lookup(content)

                # 获取这个文件的 TextNumbering 编号映射
                text_number_mapping = None
                if number_mappings:
                    text_number_mapping = number_mappings.get(filename)

                # 智能识别段落（区分采访人和受访人）
                paragraphs = self.identify_interview_paragraphs(content, filename)

                # 提取受访人有意义的句子
                respondent_sentences = self.extract_respondent_sentences(
                    paragraphs,
                    filename,
                    sentence_number_lookup=sentence_number_lookup,
                    file_path=file_path,
                    text_number_mapping=text_number_mapping,  # ← 传递编号映射
                )

                # 建立文件到句子的映射
                file_sentence_mapping[filename] = {
                    'file_path': file_path,
                    'sentences': respondent_sentences,
                    'original_content': content,
                    'paragraphs': paragraphs
                }

                # 合并文本（用于编码生成）
                all_texts.append(content)

                logger.info(f"处理文件 {filename}: 识别 {len(paragraphs)} 个段落，提取 {len(respondent_sentences)} 个受访人句子")

            except Exception as e:
                logger.error(f"处理文件 {file_path} 失败: {e}")
                continue

        # 合并所有文本
        combined_text = "\n\n".join(all_texts)

        return {
            'combined_text': combined_text,
            'file_sentence_mapping': file_sentence_mapping,
            'total_files': len(file_paths),
            'total_sentences': sum(len(data['sentences']) for data in file_sentence_mapping.values())
        }

    def identify_interview_paragraphs(self, content: str, filename: str, clean: bool = True) -> List[Dict[str, Any]]:
        """智能识别采访段落，区分采访人和受访人（支持精准提取）"""

        # 如果启用了精准提取，使用 SpeakerRoleExtractor
        if self.use_advanced_extraction and self.speaker_extractor:
            return self._identify_paragraphs_advanced(content, filename, clean)

        return self._identify_paragraphs_simple(content, filename)


    def _identify_paragraphs_advanced(self, content: str, filename: str, clean: bool = True) -> List[Dict[str, Any]]:
        """使用 SpeakerRoleExtractor 进行精准提取——保留采访者段落作为边界。"""
        try:
            all_segments = self.speaker_extractor.extract_all_segments(content, clean=clean)

            paragraphs = []
            for i, seg in enumerate(all_segments):
                role = seg.get('role', 'unknown')
                speaker_label = seg.get('speaker_label', '')
                if speaker_label:
                    speaker = speaker_label
                elif role == 'interviewer':
                    speaker = 'interviewer'
                else:
                    speaker = 'respondent'

                paragraphs.append({
                    'speaker': speaker,
                    'content': seg['text'],
                    'start_line': i,
                    'end_line': i + 1,
                    'filename': filename,
                    'confidence': seg.get('confidence', 1.0),
                    'method': seg.get('method', 'advanced')
                })

            logger.info(f"精准提取: {filename} - {len(paragraphs)} 个段落")
            return paragraphs

        except Exception as e:
            logger.error(f"精准提取失败，降级到原有逻辑: {e}")
            # 降级到原有逻辑
            self.use_advanced_extraction = False
            return self._identify_paragraphs_simple(content, filename)
    
    def _identify_paragraphs_simple(self, content: str, filename: str) -> List[Dict[str, Any]]:
        """原有的简单段落识别（兼容）"""
        paragraphs = []
        lines = content.split('\n')

        current_paragraph = []
        current_speaker = None
        paragraph_start_line = 0

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                if current_paragraph and current_speaker:
                    # 保存当前段落
                    paragraph_text = '\n'.join(current_paragraph)
                    paragraphs.append({
                        'speaker': current_speaker,
                        'content': paragraph_text,
                        'start_line': paragraph_start_line,
                        'end_line': i,
                        'filename': filename
                    })
                    current_paragraph = []
                    current_speaker = None
                continue

            # 仅根据行首明确标记切换说话人；无标记续行继承当前说话人
            speaker = self.detect_explicit_speaker(line)

            if speaker:
                if current_paragraph and current_speaker:
                    paragraph_text = '\n'.join(current_paragraph)
                    paragraphs.append({
                        'speaker': current_speaker,
                        'content': paragraph_text,
                        'start_line': paragraph_start_line,
                        'end_line': i,
                        'filename': filename
                    })
                current_paragraph = [line]
                current_speaker = speaker
                paragraph_start_line = i
            elif current_speaker:
                current_paragraph.append(line)

        # 添加最后一个段落
        if current_paragraph and current_speaker:
            paragraph_text = '\n'.join(current_paragraph)
            paragraphs.append({
                'speaker': current_speaker,
                'content': paragraph_text,
                'start_line': paragraph_start_line,
                'end_line': len(lines),
                'filename': filename
            })

        return paragraphs

    def detect_explicit_speaker(self, line: str) -> Optional[str]:
        """仅根据行首明确说话人标记检测，续行无标记时返回 None（由段落逻辑继承）。"""
        interviewer_patterns = [
            r'^[问Qq][：:]', r'^[Aa][：:]', r'^提问[：:]', r'^采访[：:]', r'^访谈[：:]',
            r'^主持人[：:]', r'^记者[：:]', r'^访员[：:]', r'^采访者[：:]',
        ]
        respondent_patterns = [
            r'^[答][：:]', r'^[Bb][：:]', r'^回答[：:]', r'^受访[：:]', r'^被访[：:]',
            r'^嘉宾[：:]', r'^专家[：:]', r'^受访者[：:]',
        ]

        for pattern in interviewer_patterns:
            if re.search(pattern, line):
                return "interviewer"

        for pattern in respondent_patterns:
            if re.search(pattern, line):
                return "respondent"

        return None

    def detect_speaker(self, line: str) -> Optional[str]:
        """检测说话人（含内容启发式，供其他模块使用；段落分组请用 detect_explicit_speaker）。"""
        explicit = self.detect_explicit_speaker(line)
        if explicit:
            return explicit

        if self.is_interviewer_line(line):
            return "interviewer"
        if self.is_respondent_line(line):
            return "respondent"

        return None

    def is_interviewer_line(self, line: str) -> bool:
        """判断是否是采访人说的话"""
        interviewer_indicators = [
            '请问', '想问', '了解', '采访', '提问', '访谈', '什么', '为什么', '如何', '怎样',
            '吗？', '呢？', '负责什么', '负责哪块', '收益如何', '感觉如何', '如何改进',
            '有什么困难', '什么挑战', '未来计划', '目标是什么'
        ]

        # 包含疑问词的短句很可能是采访人
        if any(indicator in line for indicator in interviewer_indicators):
            return True

        # 以问号结尾的短句
        if line.endswith(('？', '?')) and len(line) < 50:
            return True

        return False

    def is_respondent_line(self, line: str) -> bool:
        """判断是否是受访人说的话"""
        # 排除明显的采访人特征
        if self.is_interviewer_line(line):
            return False

        respondent_indicators = [
            '我们', '我的', '我觉得', '我认为', '我们的', '团队', '工作', '项目',
            '负责', '管理', '开发', '设计', '测试', '经验', '感受', '收获',
            '成果', '效果', '改进', '优化', '提升', '解决', '处理', '应对'
        ]

        # 包含工作内容或个人感受的较长的句子
        content_indicators = any(indicator in line for indicator in respondent_indicators)

        # 长度判断：受访人的回答通常较长
        if len(line) > 25 and content_indicators:
            return True

        return False  # 不再仅凭长度判定，避免采访者续行被误判

    def _normalize_for_sentence_lookup(self, text: str) -> str:
        normalized = re.sub(r'\s*\[[A-Z]?\d+\]', '', str(text or ''))
        normalized = re.sub(r'\s+', '', normalized)
        return normalized.strip()

    def _build_sentence_number_lookup(self, text: str) -> List[Tuple[int, str]]:
        parts = re.split(r'([\u3002\uFF01\uFF1F!? \n\r])', str(text or ''))
        sentences = []
        for i in range(0, len(parts) - 1, 2):
            sentence = parts[i]
            if i + 1 < len(parts):
                sentence += parts[i + 1]
            sentence = sentence.strip()
            if sentence:
                sentences.append(sentence)
        if len(parts) % 2 == 1:
            last_part = parts[-1].strip()
            if last_part:
                if sentences:
                    sentences[-1] += last_part
                else:
                    sentences.append(last_part)

        lookup = []
        for sentence in sentences:
            self._processing_sentence_counter = getattr(self, '_processing_sentence_counter', 0) + 1
            lookup.append((self._processing_sentence_counter, self._normalize_for_sentence_lookup(sentence)))
        return lookup

    def _lookup_sentence_number(self, sentence: str, sentence_number_lookup: List[Tuple[int, str]]) -> str:
        target = self._normalize_for_sentence_lookup(sentence)
        if not target:
            return ''
        for number, numbered_sentence in sentence_number_lookup or []:
            if target == numbered_sentence:
                return str(number)
        for number, numbered_sentence in sentence_number_lookup or []:
            if len(target) >= 8 and (target in numbered_sentence or numbered_sentence in target):
                shorter = min(len(target), len(numbered_sentence))
                longer = max(len(target), len(numbered_sentence))
                overlap_ratio = shorter / max(1, longer)
                if overlap_ratio >= 0.75:
                    return str(number)
        return ''

    def _merge_consecutive_same_speaker(self, paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Merge consecutive paragraphs with identical speaker labels.

        Content is joined by '。'. Metadata is taken from the first paragraph
        in each merged run; end_line is updated from the last paragraph.
        """
        if not paragraphs:
            return []

        merged = []
        current = dict(paragraphs[0])  # shallow copy

        for para in paragraphs[1:]:
            if para['speaker'] == current['speaker']:
                # Merge content
                current['content'] = current['content'].rstrip('。') + '。' + para['content']
                current['end_line'] = para.get('end_line', current.get('end_line', 0))
            else:
                merged.append(current)
                current = dict(para)

        merged.append(current)
        return merged

    def _split_by_length(self, text: str, max_len: int = 500) -> List[str]:
        """Split text into chunks not exceeding max_len characters.

        Prefers splitting at sentence-ending punctuation (。！？!?),
        falls back to clause breaks (，,；;), then hard-splits at max_len.
        """
        if len(text) <= max_len:
            return [text]

        chunks = []
        remaining = text

        while len(remaining) > max_len:
            # Search for the last natural break within max_len
            window = remaining[:max_len]
            split_at = None

            # Priority 1: last sentence-ending punctuation before max_len
            for m in re.finditer(r'[。！？!?]', window):
                split_at = m.end()

            # Priority 2: last clause break before max_len
            if split_at is None:
                for m in re.finditer(r'[，,；;]', window):
                    split_at = m.end()

            # Priority 3: hard split at max_len
            if split_at is None or split_at == 0:
                split_at = max_len

            chunks.append(remaining[:split_at])
            remaining = remaining[split_at:]

        if remaining:
            chunks.append(remaining)

        return chunks

    def _clean_paragraph_content(self, content: str) -> str:
        """清理段落内容：去除说话人标记、时间戳、特殊符号等。"""
        if not content:
            return ""
        # 1. 清理标准标记
        content = content.replace("受访者：", "").replace("采访者：", "")
        content = content.replace("受访者:", "").replace("采访者:", "")
        # 2. 清理所有类型的说话人标记（行首，使用MULTILINE匹配多行文本）
        content = re.sub(r'^[一-龥]+\d+[：:]\s*', '', content, flags=re.MULTILINE)
        content = re.sub(r'^[一-龥]+\d+\s*', '', content, flags=re.MULTILINE)
        # 3. 清理简写标记（问：、答：、Q:、A:等）
        content = re.sub(r'^(问|答|Q|A)[：:]\s*', '', content, flags=re.MULTILINE)
        content = content.strip()
        # 4. 清理特殊符号
        content = content.replace('●', '').replace('○', '').replace('◆', '').replace('◇', '')
        content = content.replace('■', '').replace('□', '').replace('▲', '').replace('△', '')
        # 5. 清理时间戳
        content = re.sub(r'\d{2}:\d{2}(?::\d{2})?', '', content)
        # 6. 清理多余的空白
        content = re.sub(r'\s+', ' ', content)
        return content.strip()

    def get_speaker_block_sentences(self, content: str, filename: str, clean: bool = True) -> List[Dict[str, Any]]:
        """返回所有说话人块句子（含采访者+受访者），使用与extract_respondent_sentences相同的合并和切分逻辑。

        Args:
            content: 原始文本
            filename: 文件名
            clean: 是否清洗说话人标签/时间戳等。False 时保留原文，仅加编号。

        Returns:
            List[Dict]: [{'speaker': ..., 'content': ...}, ...]
        """
        paragraphs = self.identify_interview_paragraphs(content, filename, clean=clean)
        if not paragraphs:
            return []

        split_mode = getattr(Config, 'SENTENCE_SPLIT_MODE', 'punctuation')
        max_len = getattr(Config, 'SENTENCE_MAX_LENGTH', 500)

        if split_mode == 'speaker_block':
            paragraphs = self._merge_consecutive_same_speaker(paragraphs)

        result = []
        for para in paragraphs:
            text = self._clean_paragraph_content(para['content']) if clean else para['content']
            if not text:
                continue

            if split_mode == 'speaker_block':
                chunks = self._split_by_length(text, max_len)
            else:
                chunks = self.split_into_sentences(text)

            for chunk in chunks:
                chunk = chunk.strip()
                if chunk:
                    result.append({'speaker': para['speaker'], 'content': chunk})

        return result

    def extract_respondent_sentences(
        self,
        paragraphs: List[Dict[str, Any]],
        filename: str,
        sentence_number_lookup: Optional[List[Tuple[int, str]]] = None,
        file_path: Optional[str] = None,
        text_number_mapping: Optional[Dict[int, str]] = None,
    ) -> List[Dict[str, Any]]:
        """从受访人段落中提取有意义的句子"""
        respondent_sentences = []

        # ── speaker_block 模式：合并连续同说话人段落 ──
        split_mode = getattr(Config, 'SENTENCE_SPLIT_MODE', 'punctuation')
        sentence_max_len = getattr(Config, 'SENTENCE_MAX_LENGTH', 500)

        if split_mode == 'speaker_block':
            paragraphs = self._merge_consecutive_same_speaker(paragraphs)

        for paragraph in paragraphs:
            if paragraph['speaker'] != 'interviewer':
                content = paragraph['content']
                
                # 清理角色标记（关键！）
                # 1. 清理标准标记
                content = content.replace("受访者：", "").replace("采访者：", "")
                content = content.replace("受访者:", "").replace("采访者:", "")
                
                # 2. 清理所有类型的说话人标记（行首，使用MULTILINE匹配多行文本）
                # 匹配模式：说话人1：、说话人2:、里弄管家3：、受访者4: 等
                # 先清理带冒号的完整标记
                content = re.sub(r'^[\u4e00-\u9fa5]+\d+[：:]\s*', '', content, flags=re.MULTILINE)
                # 再清理不带冒号的情况（容错处理）
                content = re.sub(r'^[\u4e00-\u9fa5]+\d+\s*', '', content, flags=re.MULTILINE)
                
                # 3. 清理简写标记（问：、答：、Q:、A:等）
                content = re.sub(r'^(问|答|Q|A)[：:]\s*', '', content, flags=re.MULTILINE)
                
                content = content.strip()
                
                # 4. 清理特殊符号（●○◆◇■□▲△等）
                content = content.replace('●', '').replace('○', '').replace('◆', '').replace('◇', '')
                content = content.replace('■', '').replace('□', '').replace('▲', '').replace('△', '')
                
                # 5. 清理时间戳（如 00:14, 01:23:45）
                content = re.sub(r'\d{2}:\d{2}(?::\d{2})?', '', content)
                
                # 6. 清理多余的空白
                content = re.sub(r'\s+', ' ', content)
                
                content = content.strip()
                
                if split_mode == 'speaker_block':
                    sentences = self._split_by_length(content, sentence_max_len)
                else:
                    sentences = self.split_into_sentences(content)

                for sentence in sentences:
                    if self.is_meaningful_sentence(sentence):
                        original_sentence = sentence  # 保存原始文本（未清理）
                        # 移除一阶编码标记（如 [A1], [A2] 等）
                        clean_sentence = re.sub(r'\s*\[A\d+\]', '', sentence)
                        clean_sentence = clean_sentence.strip()
                        
                        if clean_sentence and len(clean_sentence) >= 5:
                            # 查找对应的 TextNumbering 编号
                            text_number = None
                            numbered_sentence = clean_sentence
                            
                            # 策略1：优先从 original_sentence 中提取编号（如果文本已包含[数字]标记）
                            if text_number is None:
                                marker_match = re.search(r'\[(\d+)\]', original_sentence)
                                if marker_match:
                                    text_number = int(marker_match.group(1))
                                    numbered_sentence = f"{clean_sentence} [{text_number}]"
                            
                            # 策略2：通过 text_number_mapping 查找编号
                            if text_number is None and text_number_mapping:
                                text_number = self._find_text_number(original_sentence, text_number_mapping)
                                if text_number:
                                    numbered_sentence = f"{clean_sentence} [{text_number}]"
                            
                            # sentence_id 统一为全文 TextNumbering 主键；仅在缺失时回退到本地 lookup
                            stable_sentence_id = None
                            if text_number is not None:
                                stable_sentence_id = str(text_number)
                            else:
                                lookup_sentence_id = self._lookup_sentence_number(clean_sentence, sentence_number_lookup or [])
                                if lookup_sentence_id:
                                    stable_sentence_id = lookup_sentence_id
                            
                            # 关键修复：content字段使用带编号的文本，确保编码生成器能从中提取编号
                            # 当text_number查找失败时，仍能通过content字段提取编号
                            sentence_info = {
                                'content': numbered_sentence,  # 使用带编号的文本
                                'original_content': original_sentence,  # 保存清理后的内容
                                'paragraph_content': content[:100] + '...' if len(content) > 100 else content,
                                'filename': filename,
                                'speaker': 'respondent',
                                'start_position': 0,
                                'end_position': len(clean_sentence),
                                'text_number': text_number,  # TextNumbering 编号
                                'numbered_sentence': numbered_sentence,  # 带编号的句子
                            }
                            if file_path:
                                sentence_info['file_path'] = file_path
                            if stable_sentence_id:
                                sentence_info['sentence_id'] = stable_sentence_id
                                sentence_info['code_id'] = stable_sentence_id
                            respondent_sentences.append(sentence_info)

        return respondent_sentences
    def split_into_sentences(self, text: str) -> List[str]:
        """将文本分割成句子，保留句末标点"""
        parts = re.split(r'([。！？!?])', text)
        sentences = []
        for i in range(0, len(parts) - 1, 2):
            sentence = (parts[i] + (parts[i + 1] if i + 1 < len(parts) else '')).strip()
            if sentence and len(sentence) >= 5:
                sentences.append(sentence)
        if len(parts) % 2 == 1:
            last = parts[-1].strip()
            if last and len(last) >= 5:
                sentences.append(last)
        return sentences

    def is_meaningful_sentence(self, sentence: str) -> bool:
        """判断句子是否有意义"""
        # 过滤无意义短句
        meaningless_patterns = [
            r'^为什么[？?]?$', r'^我不知道[。.]?$', r'^什么意思[？?]?$',
            r'^然后呢[？?]?$', r'^还有吗[？?]?$', r'^嗯+$', r'^啊+$',
            r'^对+$', r'^是+$', r'^好+$', r'^行+$'
        ]

        for pattern in meaningless_patterns:
            if re.match(pattern, sentence.strip()):
                return False

        # 检查句子长度和内容
        if len(sentence.strip()) < 8:
            return False

        # 检查是否包含实质性内容
        meaningful_keywords = [
            '因为', '所以', '但是', '然而', '因此', '于是', '然后',
            '工作', '团队', '管理', '发展', '创新', '问题', '解决',
            '感觉', '认为', '觉得', '应该', '需要', '重要', '负责',
            '方法', '技术', '项目', '质量', '检测', '领导', '变革',
            '主要', '关键', '重点', '核心', '特别', '尤其', '总之',
            '目标', '成果', '效果', '影响', '原因', '结果', '过程'
        ]

        return any(keyword in sentence for keyword in meaningful_keywords) or len(sentence) >= 15

    def clean_text(self, text: str) -> str:
        """清洗文本 - 增强版，过滤录音转录信息"""
        if not text:
            return ""

        # 按行处理
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 过滤说话人标记（如：说话人1 00:51, Speaker1 12:30等）
            if self.is_speaker_time_mark(line):
                continue

            # 过滤纯时间标记 (00:01, 12:30:45)
            if re.match(r'^\d{1,2}:\d{2}(:\d{2})?$', line):
                continue

            # 过滤提问人标记（如：采访者：，提问者：，Interviewer: 等）
            if self.is_interviewer_mark(line):
                continue

            # 清理多余的标点符号
            line = re.sub(r'[，,；;]{2,}', '，', line)  # 多个逗号/分号合并为一个
            line = re.sub(r'[。！？!?]{2,}', '。', line)  # 多个句号合并为一个

            cleaned_lines.append(line)

        # 重新组合文本
        cleaned_text = '\n'.join(cleaned_lines)

        # 进一步清理整个文本
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)  # 合并多个空格
        cleaned_text = re.sub(r'([。！？!?])\s*', r'\1\n', cleaned_text)  # 在句子结束符后换行

        return cleaned_text.strip()

    def is_speaker_time_mark(self, line: str) -> bool:
        """判断是否是说话人时间标记"""
        # 匹配模式：说话人 + 时间
        patterns = [
            r'^[a-zA-Z\u4e00-\u9fa5]+\s*\d+\s*\d+:\d+',  # 说话人1 00:51
            r'^[a-zA-Z\u4e00-\u9fa5]+\d*\s*\d+:\d+',  # Speaker1 12:30
            r'^[a-zA-Z\u4e00-\u9fa5]+\s*\d+:\d+',  # 采访者 01:23
            r'^[a-zA-Z]+\s*\d+:\d+',  # Interviewer 12:30
            r'^[\u4e00-\u9fa5]+\s*\d+:\d+',  # 受访者 02:15
        ]

        for pattern in patterns:
            if re.match(pattern, line):
                return True
        return False

    def is_interviewer_mark(self, line: str) -> bool:
        """判断是否是采访人标记"""
        # 匹配采访人标记
        interviewer_marks = [
            r'^[a-zA-Z\u4e00-\u9fa5]+[：:]',  # 任何中文或英文后跟冒号
        ]

        # 但排除可能是受访人内容的情况
        respondent_indicators = ['我们', '我的', '我觉得', '我认为', '我们的', '负责', '管理', '开发']

        for pattern in interviewer_marks:
            if re.match(pattern, line):
                # 检查是否包含受访人特征，如果是则保留
                content_after_colon = re.sub(r'^[a-zA-Z\u4e00-\u9fa5]+[：:]', '', line).strip()
                if any(indicator in content_after_colon for indicator in respondent_indicators):
                    return False
                return True
        return False



    def detect_speaker_enhanced(self, line: str) -> Optional[str]:
        """增强的说话人检测 - 过滤录音转录信息"""
        # 首先过滤掉明显的时间标记等
        if self.is_speaker_time_mark(line):
            return None

        # 采访人模式 - 增强
        interviewer_patterns = [
            r'^[问Qq][：:]', r'^提问[：:]', r'^采访[：:]', r'^访谈[：:]',
            r'^主持人[：:]', r'^记者[：:]', r'^访员[：:]', r'^请问',
            r'^您觉得', r'^您认为', r'^你们', r'^公司', r'^团队',
            r'^Interviewer[：:]', r'^Host[：:]', r'^Reporter[：:]'
        ]

        # 受访人模式 - 增强
        respondent_patterns = [
            r'^[答Aa][：:]', r'^回答[：:]', r'^受访[：:]', r'^被访[：:]',
            r'^嘉宾[：:]', r'^专家[：:]', r'^我们', r'^我的', r'^我觉得',
            r'^我认为', r'^我们的', r'^负责', r'^管理', r'^开发',
            r'^Interviewee[：:]', r'^Guest[：:]', r'^Expert[：:]'
        ]

        # 检查采访人模式
        for pattern in interviewer_patterns:
            if re.search(pattern, line):
                return "interviewer"

        # 检查受访人模式
        for pattern in respondent_patterns:
            if re.search(pattern, line):
                return "respondent"

        # 基于内容判断
        if self.is_interviewer_line_enhanced(line):
            return "interviewer"
        elif self.is_respondent_line_enhanced(line):
            return "respondent"

        return None

    def is_interviewer_line_enhanced(self, line: str) -> bool:
        """判断是否是采访人说的话 - 增强版"""
        # 过滤掉短的时间标记等
        if len(line) < 10 and (re.match(r'^\d+:\d+', line) or re.match(r'^[a-zA-Z\u4e00-\u9fa5]+\d*$', line)):
            return False

        interviewer_indicators = [
            '请问', '想问', '了解', '采访', '提问', '访谈', '什么', '为什么', '如何', '怎样',
            '吗？', '呢？', '负责什么', '负责哪块', '收益如何', '感觉如何', '如何改进',
            '有什么困难', '什么挑战', '未来计划', '目标是什么', '能不能', '是否可以'
        ]

        # 包含疑问词的短句很可能是采访人
        if any(indicator in line for indicator in interviewer_indicators):
            return True

        # 以问号结尾的短句
        if line.endswith(('？', '?')) and len(line) < 50:
            return True

        return False

    def is_respondent_line_enhanced(self, line: str) -> bool:
        """判断是否是受访人说的话 - 增强版"""
        # 排除明显的采访人特征
        if self.is_interviewer_line_enhanced(line):
            return False

        # 过滤掉短的时间标记等
        if len(line) < 10 and (re.match(r'^\d+:\d+', line) or re.match(r'^[a-zA-Z\u4e00-\u9fa5]+\d*$', line)):
            return False

        respondent_indicators = [
            '我们', '我的', '我觉得', '我认为', '我们的', '团队', '工作', '项目',
            '负责', '管理', '开发', '设计', '测试', '经验', '感受', '收获',
            '成果', '效果', '改进', '优化', '提升', '解决', '处理', '应对'
        ]

        # 包含工作内容或个人感受的较长的句子
        content_indicators = any(indicator in line for indicator in respondent_indicators)

        # 长度判断：受访人的回答通常较长
        if len(line) > 25 and content_indicators:
            return True

        return False  # 不再仅凭长度判定，避免采访者续行被误判

    def merge_coding_data(self, standard_answers: Dict[str, Any], current_codes: Dict[str, Any]) -> Dict[str, Any]:
        """合并标准答案和当前编码数据"""
        try:
            # 深拷贝标准答案，避免修改原数据
            merged = copy.deepcopy(standard_answers)

            logger.info(f"开始合并编码数据: 标准答案有{len(standard_answers)}个三阶编码")

            for third_cat, second_cats in current_codes.items():
                # 清理三阶编码名称（移除编号前缀）
                clean_third_cat = self.clean_category_name(third_cat)

                if clean_third_cat not in merged:
                    merged[clean_third_cat] = {}
                    logger.info(f"新增三阶编码: {clean_third_cat}")

                for second_cat, first_contents in second_cats.items():
                    # 清理二阶编码名称（移除编号前缀）
                    clean_second_cat = self.clean_category_name(second_cat)

                    if clean_second_cat not in merged[clean_third_cat]:
                        merged[clean_third_cat][clean_second_cat] = []
                        logger.info(f"新增二阶编码: {clean_third_cat} -> {clean_second_cat}")

                    # 处理一阶编码内容
                    existing_first_contents = set(merged[clean_third_cat][clean_second_cat])

                    for content_data in first_contents:
                        if isinstance(content_data, dict):
                            # 提取一阶编码内容（清理编号前缀）
                            if "numbered_content" in content_data:
                                first_content = self.clean_first_level_content(content_data["numbered_content"])
                            elif "content" in content_data:
                                first_content = self.clean_first_level_content(content_data["content"])
                            else:
                                first_content = str(content_data)
                        else:
                            first_content = self.clean_first_level_content(str(content_data))

                        # 只有在新内容时才添加
                        if first_content and first_content not in existing_first_contents:
                            merged[clean_third_cat][clean_second_cat].append(first_content)
                            existing_first_contents.add(first_content)
                            logger.debug(f"新增一阶编码: {clean_third_cat} -> {clean_second_cat} -> {first_content[:30]}...")

            # 统计合并结果
            total_third = len(merged)
            total_second = sum(len(categories) for categories in merged.values())
            total_first = sum(
                len(first_contents)
                for categories in merged.values()
                for first_contents in categories.values()
            )

            logger.info(f"合并完成: {total_third}三阶, {total_second}二阶, {total_first}一阶")
            return merged

        except Exception as e:
            logger.error(f"合并编码数据失败: {e}")
            return standard_answers

    def clean_category_name(self, category_name: str) -> str:
        """清理类别名称，移除编号前缀"""
        cleaned = re.sub(r'^[A-Z]\d*\s*', '', category_name.strip())
        return cleaned

    def clean_first_level_content(self, content: str) -> str:
        """清理一阶编码内容，移除编号前缀"""
        cleaned = re.sub(r'^[A-Z]\d+\s*', '', content.strip())
        return cleaned

    def export_structured_codes_to_table(self, file_path: str, structured_codes: Dict[str, Any],
                                          higher_level_data: list = None) -> bool:
        """导出编码结构为表格格式（支持多阶编码）"""
        if higher_level_data is None:
            higher_level_data = []
        try:
            table_data = []
            covered = set()

            has_higher = len(higher_level_data) > 0

            if has_higher:
                self._extract_table_rows_from_higher(higher_level_data, [], table_data, covered)

            for third_category, second_categories in structured_codes.items():
                clean_third = self.clean_category_name(third_category)
                for second_category, first_contents in second_categories.items():
                    clean_second = self.clean_category_name(second_category)
                    for content_data in first_contents:
                        if isinstance(content_data, dict):
                            first_level_content = content_data.get('content', '')
                            first_level_content = re.sub(r'^[A-Z]\d+\s*', '', first_level_content).strip()
                            code_id = content_data.get('code_id', '')
                            key = (clean_third, clean_second, code_id)
                        else:
                            first_level_content = str(content_data)
                            key = (clean_third, clean_second, first_level_content)

                        if key not in covered:
                            row = {
                                "六阶编码": "",
                                "五阶编码": "",
                                "四阶编码": "",
                                "三阶编码": clean_third,
                                "二阶编码": clean_second,
                                "一阶编码": first_level_content
                            }
                            table_data.append(row)

            headers = ["六阶编码", "五阶编码", "四阶编码", "三阶编码", "二阶编码", "一阶编码"] if has_higher else ["三阶编码", "二阶编码", "一阶编码"]

            df = pd.DataFrame(table_data, columns=headers) if table_data else pd.DataFrame(columns=headers)

            if file_path.endswith('.xlsx'):
                df.to_excel(file_path, index=False, engine='openpyxl')
            elif file_path.endswith('.csv'):
                df.to_csv(file_path, index=False, encoding='utf-8-sig')
            else:
                file_path = file_path.replace('.json', '.xlsx')
                df.to_excel(file_path, index=False, engine='openpyxl')

            logger.info(f"表格格式编码已导出: {len(table_data)} 行数据")
            return True

        except Exception as e:
            logger.error(f"导出表格格式失败: {e}")
            return False

    def _extract_table_rows_from_higher(self, items: list, parent_path: list,
                                         table_data: list, covered: set):
        """从高阶编码数据递归提取表格行"""
        for item_data in items:
            text = item_data.get('text', '')
            data = item_data.get('data', {}) or {}
            level = data.get('level', 3) if isinstance(data, dict) else 3

            current_path = parent_path + [text]
            children = item_data.get('children', [])

            if level in (4, 5, 6):
                if children:
                    self._extract_table_rows_from_higher(children, current_path, table_data, covered)
            elif level == 3:
                for child in children:
                    child_data = child.get('data', {}) or {}
                    child_level = child_data.get('level', 2) if isinstance(child_data, dict) else 2
                    if child_level == 2:
                        second_text = child.get('text', '')
                        grand_children = child.get('children', [])
                        for gc in grand_children:
                            gc_data = gc.get('data', {}) or {}
                            gc_level = gc_data.get('level', 1) if isinstance(gc_data, dict) else 1
                            if gc_level == 1:
                                code_id = gc_data.get('code_id', '') if isinstance(gc_data, dict) else ''
                                content = gc_data.get('content', '') if isinstance(gc_data, dict) else ''
                                content = re.sub(r'^[A-Z]\d+\s*', '', content).strip() if content else gc.get('text', '')
                                clean_third = self.clean_category_name(text)
                                clean_second = self.clean_category_name(second_text)
                                key = (clean_third, clean_second, code_id)
                                covered.add(key)

                                row = {
                                    "六阶编码": "",
                                    "五阶编码": "",
                                    "四阶编码": "",
                                    "三阶编码": text,
                                    "二阶编码": second_text,
                                    "一阶编码": content
                                }
                                self._fill_higher_columns(row, current_path)
                                table_data.append(row)

    def _fill_higher_columns(self, row: dict, parent_path: list):
        """填充高阶列"""
        col_names = ["六阶编码", "五阶编码", "四阶编码"]
        start_idx = 3 - len(parent_path)
        for i, name in enumerate(parent_path):
            if start_idx + i < 3:
                row[col_names[start_idx + i]] = name

    def export_for_training_format(self, file_path: str, structured_codes: Dict[str, Any]) -> bool:
        """导出为训练数据格式"""
        try:
            # 构建与标准答案相同的结构
            training_format = {}

            for third_category, second_categories in structured_codes.items():
                clean_third = self.clean_category_name(third_category)
                training_format[clean_third] = {}

                for second_category, first_contents in second_categories.items():
                    clean_second = self.clean_category_name(second_category)

                    # 提取所有一阶编码内容
                    first_level_contents = []
                    for content_data in first_contents:
                        if isinstance(content_data, dict):
                            content = content_data.get('content', '')
                            # 去掉编号前缀，只保留内容
                            content = re.sub(r'^[A-Z]\d+\s*', '', content).strip()
                            first_level_contents.append(content)
                        else:
                            first_level_contents.append(str(content_data))

                    training_format[clean_third][clean_second] = first_level_contents

            # 导出为JSON
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(training_format, f, ensure_ascii=False, indent=2)

            logger.info(f"训练格式编码已导出: {file_path}")
            return True

        except Exception as e:
            logger.error(f"导出训练格式失败: {e}")
            return False

    def get_timestamp(self) -> str:
        """获取时间戳"""
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def _find_text_number(self, text: str, text_number_mapping: Dict[int, str]) -> Optional[int]:
        """
        查找文本对应的 TextNumbering 编号（平衡版：保证样本量的同时提高准确性）
        
        Args:
            text: 句子文本
            text_number_mapping: {number: text} 映射
        
        Returns:
            编号或None
        """
        if not text_number_mapping:
            return None
        
        original_text = text
        text = str(text).strip()
        
        # 0. 先检查文本中是否已经包含编号标记（如 [2345]）
        marker_match = re.search(r"\[(\d+)\]", text)
        if marker_match:
            marker_num = int(marker_match.group(1))
            if marker_num in text_number_mapping:
                return marker_num
        
        # 1. 精确匹配
        for num, mapped_text in text_number_mapping.items():
            mapped_text_str = str(mapped_text).strip()
            if text == mapped_text_str:
                return num
        
        # 2. 去除句号后精确匹配
        text_no_punct = text.rstrip('。！？!?')
        for num, mapped_text in text_number_mapping.items():
            mapped_no_punct = str(mapped_text).rstrip('。！？!?')
            if text_no_punct == mapped_no_punct:
                return num
        
        # 3. 去除空格和标点后精确匹配
        text_clean = text.replace(' ', '').replace('\n', '').replace('\t', '').rstrip('。！？!?')
        for num, mapped_text in text_number_mapping.items():
            mapped_clean = str(mapped_text).replace(' ', '').replace('\n', '').replace('\t', '').rstrip('。！？!?')
            if text_clean == mapped_clean:
                return num
        
        # 4. 检查是否是编号映射中的子串（更宽松的匹配）
        for num, mapped_text in text_number_mapping.items():
            mapped_text_str = str(mapped_text).strip()
            if text in mapped_text_str or mapped_text_str in text:
                return num
        
        # 5. 智能相似度匹配（降低阈值以提高匹配率）
        if len(text_clean) >= 5:
            best_match = None
            best_similarity = 0.0
            
            for num, mapped_text in text_number_mapping.items():
                mapped_clean = str(mapped_text).replace(' ', '').replace('\n', '').replace('\t', '').rstrip('。！？!?')
                
                if len(mapped_clean) > 0:
                    # 计算字符集合相似度
                    common_chars = len(set(text_clean) & set(mapped_clean))
                    total_chars = max(len(text_clean), len(mapped_clean))
                    char_similarity = common_chars / total_chars
                    
                    # 计算长度相似度
                    length_ratio = min(len(text_clean), len(mapped_clean)) / max(len(text_clean), len(mapped_clean))
                    
                    # 计算子串匹配度
                    substring_match = 0.0
                    if text_clean in mapped_clean:
                        substring_match = len(text_clean) / len(mapped_clean)
                    elif mapped_clean in text_clean:
                        substring_match = len(mapped_clean) / len(text_clean)
                    
                    # 综合评分（降低阈值）
                    if substring_match > 0.5:
                        if char_similarity > 0.6 and length_ratio > 0.5:
                            similarity = (char_similarity + length_ratio + substring_match) / 3
                            if similarity > best_similarity:
                                best_similarity = similarity
                                best_match = num
                    else:
                        if char_similarity > 0.75 and length_ratio > 0.65:
                            similarity = (char_similarity + length_ratio) / 2
                            if similarity > best_similarity:
                                best_similarity = similarity
                                best_match = num
            
            if best_match is not None and best_similarity > 0.6:
                return best_match
        
        return None

