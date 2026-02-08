import os
import json
import shutil
from datetime import datetime
import logging
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class ProjectManager:
    """项目管理器 - 处理项目的保存和加载"""

    PROJECT_DIR_NAME = "projects"
    PROJECT_META_FILE = "project_meta.json"
    PROJECT_DATA_FILE = "project_data.json"

    def __init__(self):
        # 创建项目目录结构
        self.project_base_dir = os.path.join(os.getcwd(), self.PROJECT_DIR_NAME)
        os.makedirs(self.project_base_dir, exist_ok=True)

        # 创建分类文件夹
        self.main_projects_dir = os.path.join(self.project_base_dir, "主界面项目保存")
        self.manual_coding_dir = os.path.join(self.project_base_dir, "手动编码保存编码")
        self.coding_tree_dir = os.path.join(self.project_base_dir, "手动编码编码树保存")

        os.makedirs(self.main_projects_dir, exist_ok=True)
        os.makedirs(self.manual_coding_dir, exist_ok=True)
        os.makedirs(self.coding_tree_dir, exist_ok=True)

    def save_project(self, project_name: str, loaded_files: Dict[str, Any], structured_codes: Dict[str, Any]) -> bool:
        """保存项目到主界面项目保存文件夹"""
        try:
            # 创建项目文件夹（在主界面项目保存目录下）
            project_dir = os.path.join(self.main_projects_dir, project_name)
            os.makedirs(project_dir, exist_ok=True)

            # 保存项目元数据
            meta_data = {
                "project_name": project_name,
                "created_at": datetime.now().isoformat(),
                "updated_at": datetime.now().isoformat(),
                "file_count": len(loaded_files),
                "save_type": "main_project"  # 标记为主界面项目保存
            }

            # 保存项目数据
            project_data = {
                "loaded_files": loaded_files,
                "structured_codes": structured_codes
            }

            # 写入元数据
            meta_path = os.path.join(project_dir, self.PROJECT_META_FILE)
            with open(meta_path, 'w', encoding='utf-8') as f:
                json.dump(meta_data, f, ensure_ascii=False, indent=2)

            # 写入项目数据
            data_path = os.path.join(project_dir, self.PROJECT_DATA_FILE)
            with open(data_path, 'w', encoding='utf-8') as f:
                json.dump(project_data, f, ensure_ascii=False, indent=2)

            # 如果存在编号后的文本，也保存一份单独的Word文档文件
            for file_path, file_data in loaded_files.items():
                if 'numbered_content' in file_data and file_data['numbered_content']:
                    filename = os.path.splitext(os.path.basename(file_path))[0]
                    numbered_docx_path = os.path.join(project_dir, f"{filename}_numbered.docx")

                    # 创建Word文档
                    doc = Document()

                    # 设置标题
                    title_para = doc.add_heading('编号文本内容', 0)
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 添加编号文本内容
                    numbered_content = file_data['numbered_content']
                    # 按行分割文本并添加到文档
                    lines = numbered_content.split('\n')
                    for line in lines:
                        if line.strip():
                            para = doc.add_paragraph(line.strip())
                            # 设置字体大小
                            for run in para.runs:
                                run.font.size = Pt(10)
                        else:
                            # 添加空行
                            doc.add_paragraph('')

                    # 保存文档
                    doc.save(numbered_docx_path)

            logger.info(f"项目已保存: {project_name} 到 {project_dir}")
            return True

        except Exception as e:
            logger.error(f"保存项目失败: {e}")
            return False

    def load_project(self, project_name: str) -> tuple:
        """加载指定名称的项目（从主界面项目保存文件夹）"""
        try:
            project_dir = os.path.join(self.main_projects_dir, project_name)
            if not os.path.exists(project_dir):
                logger.error(f"项目不存在: {project_name}")
                return None, None

            # 读取项目数据
            data_path = os.path.join(project_dir, self.PROJECT_DATA_FILE)
            if not os.path.exists(data_path):
                logger.error(f"项目数据文件不存在: {data_path}")
                return None, None

            with open(data_path, 'r', encoding='utf-8') as f:
                project_data = json.load(f)

            loaded_files = project_data.get("loaded_files", {})
            structured_codes = project_data.get("structured_codes", {})

            logger.info(f"项目已加载: {project_name}")
            return loaded_files, structured_codes

        except Exception as e:
            logger.error(f"加载项目失败: {e}")
            return None, None

    def get_projects_list(self) -> list:
        """获取主界面项目列表"""
        projects = []
        if os.path.exists(self.main_projects_dir):
            for item in os.listdir(self.main_projects_dir):
                item_path = os.path.join(self.main_projects_dir, item)
                if os.path.isdir(item_path):
                    meta_path = os.path.join(item_path, self.PROJECT_META_FILE)
                    if os.path.exists(meta_path):
                        try:
                            with open(meta_path, 'r', encoding='utf-8') as f:
                                meta_data = json.load(f)
                            projects.append({
                                "name": item,
                                "created_at": meta_data.get("created_at", ""),
                                "updated_at": meta_data.get("updated_at", ""),
                                "file_count": meta_data.get("file_count", 0),
                                "save_type": meta_data.get("save_type", "main_project")
                            })
                        except:
                            # 如果元数据文件损坏，仍然列出项目名
                            projects.append({
                                "name": item,
                                "created_at": "",
                                "updated_at": "",
                                "file_count": 0,
                                "save_type": "main_project"
                            })

        # 按更新时间排序，最新的在前
        projects.sort(key=lambda x: x.get("updated_at", ""), reverse=True)
        return projects

    def get_latest_project(self) -> str:
        """获取最新的项目名称"""
        projects = self.get_projects_list()
        if projects:
            return projects[0]["name"]
        return ""